import streamlit as st
from docx import Document
from docx.shared import RGBColor
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import os
from pathlib import Path
from datetime import datetime
import tempfile

def get_keywords_from_sheet():
    """구글 시트에서 키워드와 사유를 가져오는 함수"""
    try:
        # 구글 시트 인증
        scope = ['https://spreadsheets.google.com/feeds',
                'https://www.googleapis.com/auth/drive']
        creds = ServiceAccountCredentials.from_json_keyfile_name(
            'D:/이채윤 파일/코딩/colab-408723-89110ae33a5b.json', 
            scope
        )
        client = gspread.authorize(creds)
        
        # 시트 열기
        sheet = client.open_by_url(
            'https://docs.google.com/spreadsheets/d/1eNCbstSMyQAA7CPvwb2qE7kZWg40B7Jf-fJ7ti0ABOE/edit?gid=0'
        ).worksheet('키워드')
        
        # 키워드와 사유/제안 가져오기
        keywords = sheet.col_values(2)[2:]  # B3부터
        reasons = sheet.col_values(3)[2:]   # C3부터
        
        # 딕셔너리로 변환
        keyword_notes = {}
        for keyword, reason in zip(keywords, reasons):
            if keyword.strip():  # 빈 셀 제외
                keyword_notes[keyword] = reason if reason else ''
                
        return keyword_notes
        
    except Exception as e:
        st.error(f"구글 시트 데이터 가져오기 실패: {str(e)}")
        return None

def convert_txt_to_docx(content):
    """txt 내용을 docx로 변환"""
    try:
        # docx 파일 생성
        doc = Document()
        
        # 텍스트 디코딩 시도
        if isinstance(content, bytes):
            # 여러 인코딩 시도
            for encoding in ['utf-8', 'cp949', 'euc-kr']:
                try:
                    text = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
        else:
            text = content
            
        if not text:
            st.error("파일을 읽을 수 없습니다.")
            return None
            
        # 문단 추가 및 글꼴 설정
        paragraph = doc.add_paragraph(text)
        for run in paragraph.runs:
            run.font.name = "맑은 고딕"
            run.font.ascii_font = "맑은 고딕"
            run.font.eastasia_font = "맑은 고딕"
        
        # ���시 파일로 저장
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            return tmp.name
            
    except Exception as e:
        st.error(f"변환 중 오류 발생: {str(e)}")
        return None

def highlight_keywords(content, keyword_notes):
    """키워드 강조"""
    try:
        # 파일 확장자 확인
        if content.type == "text/plain":
            # txt 파일을 docx로 변환
            file_content = content.getvalue()
            doc_path = convert_txt_to_docx(file_content)
            if not doc_path:
                return None
        else:
            # docx 파일 임시 저장
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(content.getvalue())
                doc_path = tmp.name

        # docx 파일 처리
        doc = Document(doc_path)
        
        # 모든 단락을 순회
        for paragraph in doc.paragraphs:
            # 단락의 텍스트 저장
            text = paragraph.text
            
            # 키워드 위치 찾기
            positions = []
            for keyword in keyword_notes.keys():
                start = 0
                while True:
                    index = text.find(keyword, start)
                    if index == -1:
                        break
                    positions.append((index, index + len(keyword), keyword))
                    start = index + 1
            
            if positions:
                # 위치를 정렬
                positions.sort()
                
                # 기존 runs 제거
                for run in paragraph.runs:
                    run._element.getparent().remove(run._element)
                
                # 새로운 runs 추가
                current_pos = 0
                for start, end, keyword in positions:
                    # 키워드 전 텍스트
                    if start > current_pos:
                        run = paragraph.add_run(text[current_pos:start])
                        run.font.name = "맑은 고딕"
                    
                    # 키워드 (빨간색으로, 굵게)
                    run = paragraph.add_run(keyword)
                    run.font.name = "맑은 고딕"
                    run.font.color.rgb = RGBColor(251, 65, 65)
                    run.bold = True
                    
                    # 노트 추가
                    if keyword_notes[keyword]:
                        note_run = paragraph.add_run(f" {keyword_notes[keyword]}")
                        note_run.font.name = "맑은 고딕"
                        note_run.font.color.rgb = RGBColor(92, 179, 56)
                    
                    current_pos = end
                
                # 마지막 키워드 이후 텍스트
                if current_pos < len(text):
                    run = paragraph.add_run(text[current_pos:])
                    run.font.name = "맑은 고딕"
        
        # 결과 파일 저장
        result_path = os.path.join(tempfile.gettempdir(), "검수결과.docx")
        doc.save(result_path)
        
        # 임시 파일 삭제
        try:
            os.remove(doc_path)
        except:
            pass
            
        return result_path
        
    except Exception as e:
        st.error(f"오류 발생: {str(e)}")
        return None

def main():
    st.title("의료광고 표현 검수 시스템")
    
    # 구글 시트에서 키워드 가져오기
    keyword_notes = get_keywords_from_sheet()
    if not keyword_notes:
        st.error("키워드를 가져오지 못했습니다.")
        return
    
    # 파일 업로드
    uploaded_file = st.file_uploader("검수할 파일을 선택하세요", type=['txt', 'docx'])
    
    if uploaded_file:
        if st.button("검수 시작"):
            with st.spinner("검수 중..."):
                result_path = highlight_keywords(uploaded_file, keyword_notes)
                
                if result_path:
                    # 결과 파일 다운로드 버튼
                    with open(result_path, 'rb') as f:
                        st.download_button(
                            label="검수 결과 다운로드",
                            data=f,
                            file_name=f"검수결과_{uploaded_file.name}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    # 임시 파일 삭제
                    try:
                        os.remove(result_path)
                    except:
                        pass

if __name__ == "__main__":
    main() 